# ─── DOCUMENTS.PY — Dokument-Erstellung & Datei-Lesen ────────────────────────

import re
import os
import base64
import tempfile
from datetime import datetime
from pathlib import Path
from config import OUTPUT_DIR


# ─── FILE READING ─────────────────────────────────────────────────────────────

def read_file_content(file):
    """Liest eine hochgeladene Datei und gibt Inhalt für das LLM zurück."""
    filename = file.filename.lower()

    if filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
        data = file.read()
        b64 = base64.b64encode(data).decode('utf-8')
        ext = filename.split('.')[-1]
        mime = f"image/{'jpeg' if ext == 'jpg' else ext}"
        return {"type": "image", "b64": b64, "mime": mime}

    elif filename.endswith('.pdf'):
        try:
            import pdfplumber
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                file.save(tmp.name)
                with pdfplumber.open(tmp.name) as pdf:
                    text = "\n".join(p.extract_text() or "" for p in pdf.pages)
            os.unlink(tmp.name)
            return {"type": "text", "content": f"[PDF: {file.filename}]\n{text}"}
        except:
            return {"type": "text", "content": f"[PDF konnte nicht gelesen werden: {file.filename}]"}

    elif filename.endswith('.docx'):
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                file.save(tmp.name)
                doc = Document(tmp.name)
                text = "\n".join(p.text for p in doc.paragraphs if p.text)
            os.unlink(tmp.name)
            return {"type": "text", "content": f"[Word-Dokument: {file.filename}]\n{text}"}
        except:
            return {"type": "text", "content": f"[Word-Datei konnte nicht gelesen werden: {file.filename}]"}

    elif filename.endswith(('.xlsx', '.xls')):
        try:
            import openpyxl
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                file.save(tmp.name)
                wb = openpyxl.load_workbook(tmp.name)
                lines = []
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    lines.append(f"[Blatt: {sheet}]")
                    for row in ws.iter_rows(values_only=True):
                        lines.append(" | ".join(str(c) if c is not None else "" for c in row))
            os.unlink(tmp.name)
            return {"type": "text", "content": f"[Excel: {file.filename}]\n" + "\n".join(lines)}
        except:
            return {"type": "text", "content": f"[Excel konnte nicht gelesen werden: {file.filename}]"}

    else:
        try:
            return {"type": "text", "content": file.read().decode('utf-8', errors='ignore')}
        except:
            return {"type": "text", "content": f"[Datei konnte nicht gelesen werden: {file.filename}]"}


# ─── DOCUMENT CREATORS ────────────────────────────────────────────────────────

def create_word(data, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    if data.get("titel"):
        t = doc.add_heading(data["titel"], level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = dp.add_run(f"Erstellt: {datetime.now().strftime('%d.%m.%Y')}")
    dr.font.size      = Pt(9)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    for el in data.get("inhalt", []):
        t = el.get("typ", "absatz")
        if   t == "ueberschrift1": doc.add_heading(el.get("text", ""), level=1)
        elif t == "ueberschrift2": doc.add_heading(el.get("text", ""), level=2)
        elif t == "absatz":
            if el.get("text"): doc.add_paragraph(el["text"])
        elif t == "aufzaehlung":
            for p in el.get("punkte", []): doc.add_paragraph(p, style="List Bullet")
        elif t == "nummeriert":
            for p in el.get("punkte", []): doc.add_paragraph(p, style="List Number")
        elif t == "tabelle":
            kopf   = el.get("kopfzeile", [])
            zeilen = el.get("zeilen", [])
            if kopf or zeilen:
                cols = len(kopf) if kopf else len(zeilen[0])
                tbl  = doc.add_table(rows=1 + len(zeilen), cols=cols)
                tbl.style = "Table Grid"
                if kopf:
                    for i, txt in enumerate(kopf[:cols]):
                        tbl.rows[0].cells[i].text = str(txt)
                        for para in tbl.rows[0].cells[i].paragraphs:
                            for run in para.runs: run.bold = True
                for zi, z in enumerate(zeilen):
                    for si, v in enumerate(z[:cols]):
                        tbl.rows[zi + 1].cells[si].text = str(v)
        elif t == "trennlinie":
            doc.add_paragraph("─" * 50)

    doc.save(path)
    return path


def create_excel(data, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    wb.remove(wb.active)

    tabellen = data.get("tabellen") or []
    if not tabellen:
        tbl = {"blattname": "Tabelle1", "kopfzeile": [], "zeilen": []}
        for el in data.get("inhalt", []):
            if el.get("typ") == "tabelle":
                tbl["kopfzeile"] = el.get("kopfzeile", [])
                tbl["zeilen"]    = el.get("zeilen", [])
                break
        tabellen = [tbl]

    for td in tabellen:
        ws    = wb.create_sheet(td.get("blattname", "Tabelle1")[:31])
        start = 1
        if data.get("titel"):
            ws.cell(1, 1, data["titel"]).font = Font(bold=True, size=14)
            start = 3

        kopf   = td.get("kopfzeile", [])
        zeilen = td.get("zeilen", [])
        hfill  = PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid")
        hfont  = Font(bold=True, color="FFFFFF", size=11)

        for ci, k in enumerate(kopf, 1):
            c           = ws.cell(start, ci, k)
            c.fill      = hfill
            c.font      = hfont
            c.alignment = Alignment(horizontal="center")

        thin   = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        for zi, z in enumerate(zeilen):
            fill = PatternFill(
                start_color="F2F7FB" if zi % 2 == 0 else "FFFFFF",
                end_color  ="F2F7FB" if zi % 2 == 0 else "FFFFFF",
                fill_type  ="solid"
            )
            for si, v in enumerate(z, 1):
                c        = ws.cell(start + 1 + zi, si, v)
                c.fill   = fill
                c.border = border

        for col in ws.columns:
            maxw = max((len(str(c.value)) for c in col if c.value), default=8)
            ws.column_dimensions[col[0].column_letter].width = min(maxw + 4, 40)

    wb.save(path)
    return path


def create_text(data, path):
    lines = []
    if data.get("titel"):
        lines += [data["titel"].upper(), "=" * len(data["titel"]),
                  f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}", ""]
    for el in data.get("inhalt", []):
        t = el.get("typ", "absatz")
        if t in ("ueberschrift1", "ueberschrift2"):
            lines += ["", el.get("text", "").upper(), "-" * len(el.get("text", "")), ""]
        elif t == "absatz":
            lines += [el.get("text", ""), ""]
        elif t in ("aufzaehlung", "nummeriert"):
            for i, p in enumerate(el.get("punkte", []), 1):
                lines.append(f"  {'{}. '.format(i) if t == 'nummeriert' else '• '}{p}")
            lines.append("")
    Path(path).write_text("\n".join(lines), encoding="utf-8")
    return path


def save_document(data):
    typ  = data.get("typ", "word").lower()
    name = re.sub(r'[<>:"/\\|?*]', '_',
                  data.get("dateiname", f"dokument_{datetime.now().strftime('%Y%m%d_%H%M%S')}"))
    ext  = {"word": ".docx", "excel": ".xlsx", "text": ".txt"}.get(typ, ".docx")
    path = OUTPUT_DIR / f"{name}{ext}"
    if   typ == "word":  create_word(data, path)
    elif typ == "excel": create_excel(data, path)
    else:                create_text(data, path)
    return path, name + ext
