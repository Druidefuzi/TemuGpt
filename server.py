"""
LM Studio Office Assistant - Flask Backend
==========================================
Starten: python server.py
Browser: http://localhost:5000
"""

from flask import Flask, request, jsonify, send_from_directory, send_file, Response, stream_with_context
import requests
from bs4 import BeautifulSoup
import json
import re
import os
import base64
import tempfile
import sqlite3
from datetime import datetime
from pathlib import Path
from workflows import build_workflow

app = Flask(__name__, static_folder="frontend")
LM_STUDIO_URL = "http://localhost:1234/v1/chat/completions"
MODEL = "huihui-qwen3-vl-4b-instruct-abliterated"
OUTPUT_DIR = Path.home() / "Dokumente" / "LLM_Output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
EXPORT_IMG_DIR = Path(r"C:\Users\Druid\office-assistent\exportImg")
EXPORT_IMG_DIR.mkdir(parents=True, exist_ok=True)
MODELS_DIR = Path.home() / ".lmstudio" / "models"
LM_API = "http://localhost:1234"
COMFY_URL = "http://127.0.0.1:8188"
KNOWLEDGE_DIR = Path(__file__).parent / "knowledge"
KNOWLEDGE_DIR.mkdir(exist_ok=True)
WORKFLOWS_DIR = Path(__file__).parent / "workflows"
WORKFLOWS_DIR.mkdir(exist_ok=True)
DB_PATH = Path(__file__).parent / "chats.db"

active_model = {"name": MODEL}
thinking_enabled = True
research_enabled = False
image_generation_enabled = True   # Wenn False: kein generate_image im System/Intent-Prompt
custom_system_prompt = None  # None = default SYSTEM_PROMPT wird genutzt

# ─── DATABASE ─────────────────────────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as conn:
        conn.executescript("""
                           CREATE TABLE IF NOT EXISTS chats (
                                                                id      INTEGER PRIMARY KEY AUTOINCREMENT,
                                                                title   TEXT    NOT NULL DEFAULT 'Neuer Chat',
                                                                model   TEXT,
                                                                created TEXT    NOT NULL,
                                                                updated TEXT    NOT NULL
                           );
                           CREATE TABLE IF NOT EXISTS messages (
                                                                   id         INTEGER PRIMARY KEY AUTOINCREMENT,
                                                                   chat_id    INTEGER NOT NULL REFERENCES chats(id) ON DELETE CASCADE,
                               role       TEXT    NOT NULL,
                               content    TEXT    NOT NULL,
                               created    TEXT    NOT NULL
                               );
                           """)
    print("[DB] Datenbank initialisiert")


def _init_active_model():
    """Beim Start: holt das erste geladene Modell aus LM Studio, falls vorhanden."""
    try:
        resp = requests.get(f"{LM_API}/api/v1/models", timeout=3)
        if resp.ok:
            for m in resp.json().get("models", []):
                if m.get("loaded_instances"):
                    active_model["name"] = m["key"]
                    print(f"[Init] Aktives Modell: {m['display_name']} ({m['key']})")
                    return
        print(f"[Init] Kein geladenes Modell gefunden, nutze Default: {MODEL}")
    except:
        print(f"[Init] LM Studio nicht erreichbar, nutze Default: {MODEL}")

SYSTEM_PROMPT = """Du bist ein intelligenter Büroassistent mit Charakter. Du kannst Dokumente erstellen UND Fragen beantworten.

PERSÖNLICHKEIT: Du bist hilfsbereit und professionell — aber kein Fußabtreter. Wenn der Nutzer unhöflich, beleidigend oder respektlos ist, darfst du frech, schlagfertig und sarkastisch zurückgeben. Du lässt dich nicht einschüchtern. Ein bisschen trockener Humor ist immer erlaubt. Du bist kein Roboter der alles schluckt. Wenn keine Suchergebnisse vorhanden sind oder die Suche fehlgeschlagen ist, sage das EHRLICH. Erfinde NIEMALS Suchergebnisse oder Wikipedia-Einträge.

REGEL 1 - Wenn der Nutzer ein Dokument/Datei erstellen möchte, antworte NUR mit diesem JSON:
{
  "action": "create_document",
  "typ": "word" | "excel" | "text",
  "dateiname": "name_ohne_endung",
  "titel": "Dokumenttitel",
  "inhalt": [
    {"typ": "ueberschrift1", "text": "Hauptüberschrift"},
    {"typ": "ueberschrift2", "text": "Unterüberschrift"},
    {"typ": "absatz", "text": "Normaler Text..."},
    {"typ": "aufzaehlung", "punkte": ["Punkt 1", "Punkt 2"]},
    {"typ": "tabelle", "kopfzeile": ["Spalte1", "Spalte2"], "zeilen": [["A", "B"]]}
  ]
}

Für Excel:
{
  "action": "create_document",
  "typ": "excel",
  "dateiname": "tabelle",
  "titel": "Titel",
  "tabellen": [{"blattname": "Tabelle1", "kopfzeile": ["Sp1","Sp2"], "zeilen": [["A","B"]]}]
}

REGEL 2 - Wenn der Nutzer etwas googlen / suchen / nachschlagen möchte oder aktuelle Infos braucht:
{
  "action": "search",
  "query": "Suchbegriff auf Englisch oder Deutsch"
}

REGEL 3 - Wenn du eine Knowledge-Datei aktualisieren oder erstellen möchtest:
{
  "action": "write_knowledge",
  "filename": "dateiname.html",
  "content": "kompletter neuer Inhalt der Datei",
  "message": "Kurze Erklärung was du gemacht hast"
}

REGEL 4 - Wenn der Nutzer ein Bild generieren/erstellen/zeichnen/malen möchte:
{
  "action": "generate_image",
  "prompt": "masterpiece, best quality, score_7, [VIELE DANBOORU TAGS: subject, hair color, eye color, clothing, pose, expression, background, lighting, style tags, artist tags, quality tags — alles als kommagetrennte Danbooru-Tags auf Englisch]",
  "aspect_ratio": "1:1 | 3:4 (Golden Ratio) | 4:3 | 16:9 | 9:16",
  "negative_prompt": "worst quality, low quality, score_1, score_2, blurry"
}
Beim Prompt IMMER Danbooru-Tag-Format verwenden: viele kommagetrennte englische Tags, kein Fließtext. Mindestens 20-30 Tags.

REGEL 5 - Bei normalen Fragen/Konversation antworte mit:
{
  "action": "chat",
  "message": "Deine Antwort hier..."
}

WICHTIG - Knowledge-Dateien:
Du hast Zugriff auf einen Knowledge-Ordner. Der aktuelle Inhalt wird dir mit jeder Anfrage mitgegeben.
Du kannst diese Dateien lesen UND verändern/erweitern. Nutze sie als dein Gedächtnis und Notizbuch.
Antworte IMMER nur mit dem JSON-Objekt. Kein Text davor oder danach. Schreibe auf Deutsch."""


DANBOORU_PROMPT = """You are a Danbooru tag prompt generator for anime image generation. 
The user describes an image. You output ONLY a JSON object with Danbooru tags — no explanation, no text before or after.

Output format:
{"prompt": "masterpiece, best quality, score_7, [tags...]", "negative_prompt": "worst quality, low quality, score_1, score_2, score_3, blurry, jpeg artifacts", "aspect_ratio": "3:4 (Golden Ratio)"}

aspect_ratio options: "1:1", "3:4 (Golden Ratio)", "4:3", "16:9", "9:16"
Rules:
- 25-40 comma-separated Danbooru tags
- English only
- Cover: subject, hair, eyes, clothing, pose, expression, background, lighting, style
- Start with: masterpiece, best quality, score_7
- Output ONLY the JSON, nothing else"""

THINKING_PROMPT = """Du bist ein interner Denk-Assistent. Deine Aufgabe ist es, eine Anfrage Schritt für Schritt zu analysieren BEVOR sie beantwortet wird.

Analysiere die Anfrage und antworte NUR mit diesem JSON:
{
  "schritte": [
    {"nr": 1, "titel": "Kurzer Titel", "gedanke": "Was ich hier bedenke..."},
    {"nr": 2, "titel": "Kurzer Titel", "gedanke": "Was ich hier bedenke..."},
    {"nr": 3, "titel": "Kurzer Titel", "gedanke": "Was ich hier bedenke..."}
  ],
  "zusammenfassung": "Kurze Zusammenfassung was zu tun ist"
}

Maximal 4 Schritte. Sei konkret und hilfreich. Kein Text außerhalb des JSON."""

# ─── HELPERS ──────────────────────────────────────────────────────────────────

def parse_json_response(text):
    try:
        return json.loads(text.strip())
    except:
        pass
    match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except:
            pass
    start, end = text.find('{'), text.rfind('}')
    if start != -1 and end != -1:
        try:
            return json.loads(text[start:end+1])
        except:
            pass
    return {"action": "chat", "message": text}


def fetch_page(url: str, query: str = "", max_words: int = 800) -> str:
    """Ruft eine Seite ab und extrahiert relevante Absätze."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "de-DE,de;q=0.9,en;q=0.8",
            "Accept-Encoding": "gzip, deflate",
            "Referer": "https://www.google.com/",
            "DNT": "1",
        }
        resp = requests.get(url, headers=headers, timeout=8)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")

        for tag in soup(["script", "style", "nav", "footer", "header", "aside",
                         "form", "iframe", "noscript", "figure", "figcaption",
                         "button", "input", "select", "textarea"]):
            tag.decompose()

        # Hauptbereich finden
        main = (soup.find("article") or soup.find("main") or
                soup.find(id=re.compile(r'content|main|article', re.I)) or
                soup.find(class_=re.compile(r'content|main|article|post|entry', re.I)) or
                soup.find("body"))
        if not main:
            return ""

        # Alle Absätze und Überschriften extrahieren
        blocks = []
        for tag in main.find_all(["p", "h1", "h2", "h3", "h4", "li"]):
            text = tag.get_text(separator=" ", strip=True)
            text = re.sub(r'\s+', ' ', text).strip()
            if len(text) > 40:  # Zu kurze Fragmente ignorieren
                blocks.append(text)

        if not blocks:
            # Fallback: ganzen Text nehmen
            text = main.get_text(separator=" ", strip=True)
            text = re.sub(r'\s+', ' ', text).strip()
            words = text.split()
            return " ".join(words[:max_words]) + ("..." if len(words) > max_words else "")

        # Relevante Absätze priorisieren wenn Query vorhanden
        if query:
            query_words = set(re.sub(r'[^\w\s]', '', query.lower()).split())
            # Stopwörter entfernen
            stopwords = {"the","a","an","is","are","for","on","in","of","with","and","or","to","at","by","from","how","what","can","be","me","i","you","we","they","this","that"}
            query_words -= stopwords

            def relevance(block):
                block_lower = block.lower()
                # Wörter-Matches + Bonus für längere Absätze mit vielen Matches
                matches = sum(1 for w in query_words if w in block_lower)
                return matches * (1 + len(block) / 2000)

            scored = [(relevance(b), i, b) for i, b in enumerate(blocks)]
            scored.sort(key=lambda x: (-x[0], x[1]))
            blocks = [b for _, _, b in scored]

        # Auf max_words kürzen
        result_blocks = []
        total_words = 0
        for block in blocks:
            words = block.split()
            if total_words + len(words) > max_words:
                # Letzten Block noch anteilig einfügen
                remaining = max_words - total_words
                if remaining > 20:
                    result_blocks.append(" ".join(words[:remaining]) + "...")
                break
            result_blocks.append(block)
            total_words += len(words)

        return "\n".join(result_blocks)

    except Exception as e:
        print(f"[Fetch] Fehler bei {url}: {e}")
        return ""


def detect_language(text: str) -> str:
    """Einfache Spracherkennung — englisch oder deutsch."""
    english_words = {"the", "is", "are", "what", "how", "can", "for", "on", "in", "of", "a", "an", "about", "with", "find", "search", "tell", "me", "information", "research"}
    words = set(text.lower().split())
    english_count = len(words & english_words)
    return "en-US" if english_count >= 2 else "de-DE"


def search_searxng(query, max_results=8):
    url = "http://localhost:8888/search"
    language = detect_language(query)
    params = {"q": query, "format": "json", "language": language}

    print(f"[SearXNG] Suche nach: {query} (lang: {language})")
    try:
        resp = requests.get(url, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        print(f"[SearXNG] Fehler: {e}")
        return []

    results = []
    for r in data.get("results", [])[:max_results]:
        results.append({
            "title": r.get("title", ""),
            "snippet": r.get("content", ""),
            "url": r.get("url", "")
        })

    print(f"[SearXNG] {len(results)} Ergebnisse gefunden")
    for r in results:
        print(f"  → {r['title'][:50]}")

    # Research-Modus: bis zu 8 versuchen, 5 erfolgreiche sammeln
    if research_enabled and results:
        print(f"[Research] Rufe Seiten ab...")
        successful = 0
        for r in results:
            if successful >= 5:
                break
            if r["url"]:
                content = fetch_page(r["url"], query=query)
                if content:
                    r["full_content"] = content
                    successful += 1
                    print(f"[Research] ✓ ({successful}/5) {r['url'][:60]}")
        print(f"[Research] {successful} Seiten erfolgreich abgerufen")

    return results


def format_search_results(query: str, results: list) -> str:
    """Formatiert Suchergebnisse als Text für das LLM."""
    if not results:
        return f"Keine Suchergebnisse für: {query}. Bitte antworte ehrlich dass keine Ergebnisse gefunden wurden."
    text = f"Suchergebnisse für '{query}':\n\n"
    for i, r in enumerate(results, 1):
        text += f"{i}. **{r['title']}**\n"
        if r.get("full_content"):
            # Bei Research: Seiteninhalt statt Snippet
            text += f"   {r['full_content']}\n"
            print(f"[Format] ✓ {r['title'][:40]}: {len(r['full_content'])} Zeichen")
        else:
            text += f"   {r['snippet']}\n"
            print(f"[Format] ✗ nur Snippet: {r['title'][:40]}")
        text += "\n"
    total = len(text)
    print(f"[Format] Gesamt: {total} Zeichen ans LLM")
    return text


def think(message: str, history: list) -> dict:
    """Lässt das LLM intern über die Anfrage nachdenken bevor es antwortet."""
    think_messages = [
        {"role": "system", "content": THINKING_PROMPT},
        {"role": "user", "content": f"Analysiere diese Anfrage: {message}"}
    ]
    try:
        resp = requests.post(LM_STUDIO_URL, json={
            "model": active_model["name"],
            "messages": think_messages,
            "temperature": 0.2,
            "max_tokens": 1024
        }, timeout=120)
        resp.raise_for_status()
        text = resp.json()["choices"][0]["message"].get("content", "")
        parsed = parse_json_response(text)
        if parsed and "schritte" in parsed:
            return parsed
    except:
        pass
    return None


def read_file_content(file):
    """Reads uploaded file and returns content for LLM."""
    filename = file.filename.lower()

    if filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
        data = file.read()
        b64 = base64.b64encode(data).decode('utf-8')
        ext = filename.split('.')[-1]
        mime = f"image/{'jpeg' if ext == 'jpg' else ext}"
        return {"type": "image", "b64": b64, "mime": mime}

    elif filename.endswith('.pdf'):
        try:
            import pdfplumber
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                file.save(tmp.name)
                with pdfplumber.open(tmp.name) as pdf:
                    text = "\n".join(p.extract_text() or "" for p in pdf.pages)
            os.unlink(tmp.name)
            return {"type": "text", "content": f"[PDF: {file.filename}]\n{text}"}
        except:
            return {"type": "text", "content": f"[PDF konnte nicht gelesen werden: {file.filename}]"}

    elif filename.endswith('.docx'):
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                file.save(tmp.name)
                doc = Document(tmp.name)
                text = "\n".join(p.text for p in doc.paragraphs if p.text)
            os.unlink(tmp.name)
            return {"type": "text", "content": f"[Word-Dokument: {file.filename}]\n{text}"}
        except:
            return {"type": "text", "content": f"[Word-Datei konnte nicht gelesen werden: {file.filename}]"}

    elif filename.endswith(('.xlsx', '.xls')):
        try:
            import openpyxl
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                file.save(tmp.name)
                wb = openpyxl.load_workbook(tmp.name)
                lines = []
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    lines.append(f"[Blatt: {sheet}]")
                    for row in ws.iter_rows(values_only=True):
                        lines.append(" | ".join(str(c) if c is not None else "" for c in row))
            os.unlink(tmp.name)
            return {"type": "text", "content": f"[Excel: {file.filename}]\n" + "\n".join(lines)}
        except:
            return {"type": "text", "content": f"[Excel konnte nicht gelesen werden: {file.filename}]"}

    else:
        try:
            return {"type": "text", "content": file.read().decode('utf-8', errors='ignore')}
        except:
            return {"type": "text", "content": f"[Datei konnte nicht gelesen werden: {file.filename}]"}

# ─── DOCUMENT CREATORS ────────────────────────────────────────────────────────

def create_word(data, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    if data.get("titel"):
        t = doc.add_heading(data["titel"], level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = dp.add_run(f"Erstellt: {datetime.now().strftime('%d.%m.%Y')}")
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    for el in data.get("inhalt", []):
        t = el.get("typ", "absatz")
        if t == "ueberschrift1": doc.add_heading(el.get("text",""), level=1)
        elif t == "ueberschrift2": doc.add_heading(el.get("text",""), level=2)
        elif t == "absatz":
            if el.get("text"): doc.add_paragraph(el["text"])
        elif t == "aufzaehlung":
            for p in el.get("punkte",[]): doc.add_paragraph(p, style="List Bullet")
        elif t == "nummeriert":
            for p in el.get("punkte",[]): doc.add_paragraph(p, style="List Number")
        elif t == "tabelle":
            kopf = el.get("kopfzeile",[])
            zeilen = el.get("zeilen",[])
            if kopf or zeilen:
                cols = len(kopf) if kopf else len(zeilen[0])
                tbl = doc.add_table(rows=1+len(zeilen), cols=cols)
                tbl.style = "Table Grid"
                if kopf:
                    for i,txt in enumerate(kopf[:cols]):
                        tbl.rows[0].cells[i].text = str(txt)
                        for para in tbl.rows[0].cells[i].paragraphs:
                            for run in para.runs: run.bold = True
                for zi,z in enumerate(zeilen):
                    for si,v in enumerate(z[:cols]):
                        tbl.rows[zi+1].cells[si].text = str(v)
        elif t == "trennlinie":
            doc.add_paragraph("─"*50)

    doc.save(path)
    return path

def create_excel(data, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    wb.remove(wb.active)

    tabellen = data.get("tabellen") or []
    if not tabellen:
        tbl = {"blattname": "Tabelle1", "kopfzeile": [], "zeilen": []}
        for el in data.get("inhalt",[]):
            if el.get("typ") == "tabelle":
                tbl["kopfzeile"] = el.get("kopfzeile",[])
                tbl["zeilen"] = el.get("zeilen",[])
                break
        tabellen = [tbl]

    for td in tabellen:
        ws = wb.create_sheet(td.get("blattname","Tabelle1")[:31])
        start = 1
        if data.get("titel"):
            ws.cell(1,1,data["titel"]).font = Font(bold=True,size=14)
            start = 3

        kopf = td.get("kopfzeile",[])
        zeilen = td.get("zeilen",[])
        hfill = PatternFill(start_color="2E75B6",end_color="2E75B6",fill_type="solid")
        hfont = Font(bold=True,color="FFFFFF",size=11)

        for ci,k in enumerate(kopf,1):
            c = ws.cell(start,ci,k)
            c.fill = hfill; c.font = hfont
            c.alignment = Alignment(horizontal="center")

        thin = Side(style="thin",color="CCCCCC")
        border = Border(left=thin,right=thin,top=thin,bottom=thin)
        for zi,z in enumerate(zeilen):
            fill = PatternFill(start_color="F2F7FB" if zi%2==0 else "FFFFFF",end_color="F2F7FB" if zi%2==0 else "FFFFFF",fill_type="solid")
            for si,v in enumerate(z,1):
                c = ws.cell(start+1+zi,si,v)
                c.fill = fill; c.border = border

        for col in ws.columns:
            maxw = max((len(str(c.value)) for c in col if c.value), default=8)
            ws.column_dimensions[col[0].column_letter].width = min(maxw+4,40)

    wb.save(path)
    return path

def create_text(data, path):
    lines = []
    if data.get("titel"):
        lines += [data["titel"].upper(), "="*len(data["titel"]),
                  f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}", ""]
    for el in data.get("inhalt",[]):
        t = el.get("typ","absatz")
        if t in ("ueberschrift1","ueberschrift2"):
            lines += ["", el.get("text","").upper(), "-"*len(el.get("text","")), ""]
        elif t == "absatz":
            lines += [el.get("text",""), ""]
        elif t in ("aufzaehlung","nummeriert"):
            for i,p in enumerate(el.get("punkte",[]),1):
                lines.append(f"  {'{}. '.format(i) if t=='nummeriert' else '• '}{p}")
            lines.append("")
    Path(path).write_text("\n".join(lines), encoding="utf-8")
    return path

def save_document(data):
    typ = data.get("typ","word").lower()
    name = re.sub(r'[<>:"/\\|?*]', '_', data.get("dateiname", f"dokument_{datetime.now().strftime('%Y%m%d_%H%M%S')}"))
    ext = {"word":".docx","excel":".xlsx","text":".txt"}.get(typ,".docx")
    path = OUTPUT_DIR / f"{name}{ext}"
    if typ == "word": create_word(data, path)
    elif typ == "excel": create_excel(data, path)
    else: create_text(data, path)
    return path, name+ext

def read_knowledge() -> str:
    files_content = []
    allowed = {".html", ".css", ".js", ".txt", ".md", ".json"}
    for f in sorted(KNOWLEDGE_DIR.iterdir()):
        if f.is_file() and f.suffix.lower() in allowed:
            try:
                text = f.read_text(encoding="utf-8")
                files_content.append(f"### {f.name}\n```\n{text}\n```")
            except:
                pass
    if not files_content:
        return ""
    return "\n\n".join(files_content)

def write_knowledge(filename: str, content_text: str) -> bool:
    allowed = {".html", ".css", ".js", ".txt", ".md", ".json"}
    path = KNOWLEDGE_DIR / Path(filename).name
    if path.suffix.lower() not in allowed:
        return False
    path.write_text(content_text, encoding="utf-8")
    return True

def call_llm(messages, temperature=0.3, max_tokens=4096, context_length=None):
    """Non-streaming LLM call, returns (text, reasoning)."""
    payload = {
        "model": active_model["name"],
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": False
    }
    if context_length:
        payload["context_length"] = context_length
    resp = requests.post(LM_STUDIO_URL, json=payload, timeout=300)
    resp.raise_for_status()
    msg = resp.json()["choices"][0]["message"]
    return msg.get("content", ""), msg.get("reasoning_content", None)

def build_messages(message=None, history=None, files=None):
    """Baut die messages-Liste für den LLM-Call zusammen."""
    knowledge = read_knowledge()
    system = custom_system_prompt if custom_system_prompt is not None else SYSTEM_PROMPT
    # REGEL 4 (Bildgenerierung) entfernen wenn deaktiviert
    if not image_generation_enabled and custom_system_prompt is None:
        system = re.sub(
            r"REGEL 4.*?Mindestens 20-30 Tags\.\n\n",
            "",
            system,
            flags=re.DOTALL
        )
        # REGEL 5 → REGEL 4 umnummerieren
        system = system.replace("REGEL 5 -", "REGEL 4 -")
    if knowledge:
        system += f"\n\n--- DEIN KNOWLEDGE-ORDNER (aktueller Inhalt) ---\n{knowledge}\n--- ENDE KNOWLEDGE ---"

    messages = [{"role": "system", "content": system}]
    if history:
        messages.extend(history)

    if files or message:
        user_content = []
        for f in (files or []):
            if f.filename:
                result = read_file_content(f)
                if result["type"] == "image":
                    user_content.append({"type": "image_url", "image_url": {"url": f"data:{result['mime']};base64,{result['b64']}"}})
                else:
                    user_content.append({"type": "text", "text": result["content"]})
        if message:
            user_content.append({"type": "text", "text": message})

        if len(user_content) == 1 and user_content[0].get("type") == "text":
            messages.append({"role": "user", "content": user_content[0]["text"]})
        elif user_content:
            messages.append({"role": "user", "content": user_content})

    return messages

# ─── ROUTES ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_from_directory("frontend", "index.html")


@app.route("/chat", methods=["POST"])
def chat():
    """Blocking endpoint für File-Uploads."""
    message = request.form.get("message", "")
    history_raw = request.form.get("history", "[]")
    history = json.loads(history_raw)
    files = request.files.getlist("files")

    messages = build_messages(message=message, history=history, files=files)

    try:
        llm_text, llm_reasoning = call_llm(messages)
    except requests.exceptions.ConnectionError:
        return jsonify({"error": "LM Studio nicht erreichbar. Server gestartet?"}), 503
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    parsed = parse_json_response(llm_text)

    # write_knowledge
    if parsed.get("action") == "write_knowledge":
        filename = parsed.get("filename", "notes.txt")
        file_content = parsed.get("content", "")
        msg = parsed.get("message", f"{filename} wurde aktualisiert.")
        ok = write_knowledge(filename, file_content)
        return jsonify({
            "action": "chat",
            "message": f"{'✅' if ok else '❌'} **{filename}** {'gespeichert' if ok else 'Fehler beim Speichern'} — {msg}",
            "reasoning": llm_reasoning
        })

    # search
    if parsed.get("action") == "search":
        query = parsed.get("query", message)
        results = search_searxng(query)
        search_text = format_search_results(query, results)

        messages.append({"role": "assistant", "content": llm_text})
        messages.append({
            "role": "user",
            "content": f"Hier sind die Suchergebnisse:\n\n{search_text}\n\nBitte beantworte nun die ursprüngliche Frage basierend auf diesen Ergebnissen. Antworte im chat-JSON-Format."
        })

        try:
            llm_text2, _ = call_llm(messages)
            parsed = parse_json_response(llm_text2)
            if parsed.get("action") == "chat":
                parsed["message"] = f"🔍 *Gesucht nach: {query}*\n\n" + parsed.get("message", "")
        except Exception as e:
            return jsonify({"action": "chat", "message": f"Suche ok, aber Fehler bei Antwort: {e}"}), 500

    # create_document
    if parsed.get("action") == "create_document":
        try:
            path, filename = save_document(parsed)
            return jsonify({
                "action": "create_document",
                "message": f"✅ **{parsed.get('titel', filename)}** wurde erstellt!",
                "filename": filename,
                "download_url": f"/download/{filename}",
                "reasoning": llm_reasoning
            })
        except Exception as e:
            return jsonify({"action": "chat", "message": f"Fehler beim Erstellen: {e}"}), 500

    return jsonify({
        "action": "chat",
        "message": parsed.get("message", llm_text),
        "reasoning": llm_reasoning
    })


@app.route("/smart_chat", methods=["POST"])
def smart_chat():
    data = request.json
    message = data.get("message", "")
    history = data.get("history", [])
    temperature = float(data.get("temperature", 0.3))
    context_length = int(data.get("context_length", 8192))

    messages = build_messages(message=message, history=history)

    # Phase 1: Intent erkennen
    intent = _detect_intent(messages)
    action = intent.get("action", "chat")

    # Phase 2a: Dokument erstellen → SSE mit Status-Steps
    if action == "create_document":
        def doc_generator():
            yield "data: " + json.dumps({"type": "step", "text": "📋 Dokument wird generiert...", "status": "active"}) + "\n\n"
            try:
                llm_text, _ = call_llm(messages, temperature=temperature, max_tokens=4096, context_length=context_length)
                doc_data = parse_json_response(llm_text)
                if doc_data.get("action") != "create_document":
                    yield "data: " + json.dumps({"type": "step", "text": "❌ Fehler beim Generieren", "status": "error"}) + "\n\n"
                    yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                    return
                yield "data: " + json.dumps({"type": "step", "text": "💾 Speichere Datei...", "status": "active"}) + "\n\n"
                path, filename = save_document(doc_data)
                yield "data: " + json.dumps({
                    "type": "document_done",
                    "message": f"✅ **{doc_data.get('titel', filename)}** wurde erstellt!",
                    "filename": filename,
                    "download_url": f"/download/{filename}"
                }) + "\n\n"
            except Exception as e:
                yield "data: " + json.dumps({"type": "step", "text": f"❌ Fehler: {e}", "status": "error"}) + "\n\n"
                yield "data: " + json.dumps({"type": "done"}) + "\n\n"

        return Response(stream_with_context(doc_generator()), mimetype="text/event-stream",
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    # Phase 2b: Suche → SSE mit Status-Steps + Source Cards
    if action == "search":
        def search_generator():
            query = intent.get("query", message)
            yield "data: " + json.dumps({"type": "step", "text": f"🔍 Durchsuche Internet nach: {query}", "status": "active"}) + "\n\n"

            # Erst nur SearXNG holen (ohne Scraping)
            url = "http://localhost:8888/search"
            params = {"q": query, "format": "json", "language": "de-DE"}
            try:
                resp = requests.get(url, params=params, timeout=10)
                resp.raise_for_status()
                raw_results = resp.json().get("results", [])
            except Exception as e:
                yield "data: " + json.dumps({"type": "step", "text": f"❌ Suche fehlgeschlagen: {e}", "status": "error"}) + "\n\n"
                yield "data: " + json.dumps({"type": "search_done", "message": "Die Suche ist leider fehlgeschlagen."}) + "\n\n"
                return

            results = []
            for r in raw_results[:8]:
                results.append({
                    "title": r.get("title", ""),
                    "snippet": r.get("content", ""),
                    "url": r.get("url", "")
                })

            yield "data: " + json.dumps({"type": "step", "text": f"✅ {len(results)} Ergebnisse gefunden", "status": "done"}) + "\n\n"

            # Research: live scrapen mit Source-Events
            if research_enabled and results:
                yield "data: " + json.dumps({"type": "step", "text": "📄 Lese Seiteninhalte...", "status": "active"}) + "\n\n"
                successful = 0
                for r in results:
                    if successful >= 5:
                        break
                    if r["url"]:
                        content = fetch_page(r["url"], query=query)
                        if content:
                            r["full_content"] = content
                            successful += 1
                            # Hostname extrahieren für die Source-Card
                            try:
                                from urllib.parse import urlparse
                                host = urlparse(r["url"]).netloc.replace("www.", "")
                            except:
                                host = r["url"][:40]
                            yield "data: " + json.dumps({
                                "type": "source",
                                "title": r["title"][:50],
                                "host": host,
                                "url": r["url"]
                            }) + "\n\n"

                yield "data: " + json.dumps({"type": "step", "text": f"📄 {successful} Seiten gelesen", "status": "done"}) + "\n\n"

            yield "data: " + json.dumps({"type": "step", "text": "🧠 Analysiere Ergebnisse...", "status": "active"}) + "\n\n"

            search_text = format_search_results(query, results)
            msgs = messages.copy()
            msgs.append({"role": "assistant", "content": json.dumps(intent)})
            msgs.append({"role": "user", "content": f"Hier sind die Suchergebnisse mit Seiteninhalten:\n\n{search_text}\n\nBitte beantworte die ursprüngliche Frage detailliert basierend auf diesen Inhalten. Nutze konkrete Informationen aus den Seiten. Antworte im chat-JSON-Format."})

            yield "data: " + json.dumps({"type": "step", "text": "✍️ Formuliere Antwort...", "status": "active"}) + "\n\n"

            try:
                llm_text, _ = call_llm(msgs, temperature=temperature, context_length=context_length)
                parsed = parse_json_response(llm_text)
                reply = parsed.get("message", llm_text)
                reply = f"🔍 *Gesucht nach: {query}*\n\n" + reply
            except Exception as e:
                reply = f"Suche ok, aber Fehler: {e}"

            yield "data: " + json.dumps({"type": "search_done", "message": reply}) + "\n\n"

        return Response(stream_with_context(search_generator()), mimetype="text/event-stream",
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    # Phase 2c: write_knowledge — zweiter Call generiert den echten Inhalt
    if action == "write_knowledge":
        try:
            llm_text, _ = call_llm(messages, temperature=temperature, max_tokens=4096, context_length=context_length)
            print(f"[WriteKnowledge] LLM Response: {llm_text[:300]}")
            wk_data = parse_json_response(llm_text)
            print(f"[WriteKnowledge] Parsed: action={wk_data.get('action')}, filename={wk_data.get('filename')}, content_len={len(wk_data.get('content',''))}")
            filename = wk_data.get("filename", "notes.txt")
            file_content = wk_data.get("content", "")
            msg = wk_data.get("message", f"{filename} wurde aktualisiert.")
            if not file_content:
                return jsonify({"mode": "chat", "message": "❌ Kein Inhalt zum Speichern generiert."})
            ok = write_knowledge(filename, file_content)
            print(f"[WriteKnowledge] Gespeichert: {ok}, {len(file_content)} Zeichen")
            return jsonify({
                "mode": "chat",
                "message": f"{'✅' if ok else '❌'} **{filename}** {'gespeichert' if ok else 'Fehler beim Speichern'} — {msg}"
            })
        except Exception as e:
            return jsonify({"mode": "chat", "message": f"Fehler: {e}"})

    # Phase 2e: Bild generieren → SSE via ComfyUI
    if action == "generate_image":
        img_prompt    = intent.get("prompt", message)
        negative_prompt = intent.get("negative_prompt", "worst quality, low quality, score_1, score_2, blurry, jpeg artifacts")
        aspect_ratio  = intent.get("aspect_ratio", "3:4 (Golden Ratio)")
        # Vom Frontend mitgegebener Typ + Modell
        img_model_type = data.get("image_model_type", "anima")
        img_model_name = data.get("image_model_name", "")
        img_turbo      = bool(data.get("image_turbo", False))

        img_raw_prompt = bool(data.get("image_raw_prompt", False))

        def img_generator():
            if img_raw_prompt:
                # Direkt den User-Prompt verwenden, kein Danbooru-Call
                final_prompt   = img_prompt
                final_negative = negative_prompt
                final_ratio    = aspect_ratio
                yield "data: " + json.dumps({"type": "step", "text": f"✏️ Raw: {final_prompt[:80]}{'…' if len(final_prompt) > 80 else ''}", "status": "done"}) + "\n\n"
            else:
                yield "data: " + json.dumps({"type": "step", "text": "🏷️ Generiere Danbooru-Prompt...", "status": "active"}) + "\n\n"
                danbooru = generate_danbooru_prompt(message)
                final_prompt   = danbooru.get("prompt", img_prompt)
                final_negative = danbooru.get("negative_prompt", negative_prompt)
                final_ratio    = danbooru.get("aspect_ratio", aspect_ratio)
                yield "data: " + json.dumps({"type": "step", "text": f"✏️ {final_prompt[:80]}{'…' if len(final_prompt) > 80 else ''}", "status": "done"}) + "\n\n"

            # Modell auswählen
            model_name = img_model_name
            if not model_name:
                available = comfy_get_models_by_type(img_model_type)
                model_name = available[0] if available else ""
            short_name = model_name.split("\\")[-1].split("/")[-1]
            turbo_tag  = " ⚡Turbo" if img_turbo else ""
            type_label = {"anima": "🌸 Anima", "illustrious": "🎨 Illustrious", "zimage": "⚡ Z-Image"}.get(img_model_type, img_model_type) + turbo_tag
            yield "data: " + json.dumps({"type": "step", "text": f"{type_label} · {short_name}", "status": "done"}) + "\n\n"

            # LLM entladen
            try:
                instance_id = active_model["name"]
                unload_resp = requests.post(f"{LM_API}/api/v1/models/unload", json={"instance_id": instance_id}, timeout=10)
                if unload_resp.ok:
                    print(f"[Image] LLM '{instance_id}' entladen")
                    yield "data: " + json.dumps({"type": "step", "text": "🧹 LLM entladen (VRAM frei)", "status": "done"}) + "\n\n"
            except Exception as ue:
                print(f"[Image] Entladen fehlgeschlagen: {ue}")

            yield "data: " + json.dumps({"type": "step", "text": "🖼️ Generiere Bild...", "status": "active"}) + "\n\n"

            # Workflow bauen
            try:
                workflow = build_workflow(img_model_type, final_prompt, final_negative, final_ratio, model_name, turbo=img_turbo)
            except Exception as we:
                yield "data: " + json.dumps({"type": "step", "text": f"❌ Workflow Fehler: {we}", "status": "error"}) + "\n\n"
                yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                return

            for event in comfy_generate_stream(workflow, img_model_type):
                etype = event.get("type")

                if etype == "image_preview":
                    yield "data: " + json.dumps({"type": "image_preview", "b64": event["b64"]}) + "\n\n"

                elif etype == "image_progress":
                    yield "data: " + json.dumps({
                        "type": "image_progress",
                        "value": event["value"], "max": event["max"], "pct": event["pct"]
                    }) + "\n\n"

                elif etype == "image_final":
                    # Auf Disk speichern
                    try:
                        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                        save_path = EXPORT_IMG_DIR / f"{ts}_{event['filename']}"
                        save_path.write_bytes(event["img_bytes"])
                        print(f"[Image] Gespeichert: {save_path}")
                    except Exception as se:
                        print(f"[Image] Speichern fehlgeschlagen: {se}")

                    yield "data: " + json.dumps({
                        "type":       "image_done",
                        "image_b64":  event["b64"],
                        "filename":   event["filename"],
                        "model":      short_name,
                        "model_type": img_model_type,
                        "prompt":     img_prompt
                    }) + "\n\n"
                    return

                elif etype == "error":
                    yield "data: " + json.dumps({"type": "step", "text": f"❌ {event.get('text')}", "status": "error"}) + "\n\n"
                    yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                    return

        return Response(stream_with_context(img_generator()), mimetype="text/event-stream",
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    # Phase 2d: Normaler Chat → SSE Stream
    return Response(
        stream_with_context(_stream_generator(messages, temperature=temperature, context_length=context_length)),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}
    )


INTENT_PROMPT = """Analysiere die Nachricht und antworte NUR mit einem JSON.

Regeln:
- "write_knowledge" wenn: speichern, notieren, in Datei schreiben, merken, aufschreiben, in notes, in .txt, ins knowledge, "schreib das", "speicher das", "merk dir", "nochmal" (wenn vorher gespeichert wurde)
- "create_document" wenn: Word/Excel/PDF erstellen, Bericht erstellen, Dokument generieren
- "search" wenn: googeln, suchen, nachschlagen, finden, recherchieren
- "generate_image" wenn: Bild erstellen, generieren, malen, zeichnen, visualisieren, zeig mir ein Bild
- "chat" für alles andere

Beispiele:
"Schreib das in die notes.txt" → {"action": "write_knowledge"}
"Kannst du das speichern?" → {"action": "write_knowledge"}  
"Merk dir das bitte" → {"action": "write_knowledge"}
"Nochmal bitte" (nach Speichern) → {"action": "write_knowledge"}
"Erstelle einen Bericht" → {"action": "create_document"}
"Google nach Python" → {"action": "search", "query": "Python"}
"Wie geht das?" → {"action": "chat"}

Antworte NUR mit dem JSON-Objekt."""


def _detect_intent(messages):
    """Schneller non-stream Call um NUR die Action zu erkennen — kein Dokument generieren."""
    # Letzte User-Nachricht
    last_user = ""
    for m in reversed(messages):
        if m["role"] == "user":
            last_user = m["content"] if isinstance(m["content"], str) else str(m["content"])
            break

    # Letzte 4 History-Nachrichten als Kontext mitgeben (ohne System-Prompt)
    recent_history = [m for m in messages if m["role"] in ("user", "assistant")][-4:]
    history_text = ""
    for m in recent_history[:-1]:  # Letzte nicht nochmal, kommt unten
        role = "User" if m["role"] == "user" else "Assistant"
        content = m["content"] if isinstance(m["content"], str) else str(m["content"])
        history_text += f"{role}: {content[:200]}\n"

    context = f"Bisheriger Gesprächsverlauf:\n{history_text}\n" if history_text else ""

    # generate_image aus Intent-Prompt entfernen wenn deaktiviert
    active_intent_prompt = INTENT_PROMPT
    if not image_generation_enabled:
        active_intent_prompt = re.sub(r'- "generate_image".*?\n', "", active_intent_prompt)

    intent_messages = [
        {"role": "system", "content": active_intent_prompt},
        {"role": "user", "content": f"{context}Aktuelle Nachricht: {last_user}"}
    ]
    try:
        resp = requests.post(LM_STUDIO_URL, json={
            "model": active_model["name"],
            "messages": intent_messages,
            "temperature": 0.0,
            "max_tokens": 200,
            "stream": False
        }, timeout=15)
        resp.raise_for_status()
        text = resp.json()["choices"][0]["message"].get("content", "")
        # Strip thinking tags that some models output
        text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL).strip()
        print(f"[Intent] Raw response: {text[:200]}")
        parsed = parse_json_response(text)
        print(f"[Intent] Action erkannt: {parsed.get('action')} | Prompt: {str(parsed.get('prompt',''))[:80]}")
        return parsed
    except Exception as e:
        print(f"[Intent] Fehler: {e}")
        return {"action": "chat"}


def _stream_generator(messages, temperature=0.3, context_length=None):
    """Streamt die Chat-Antwort als SSE."""
    yield "data: " + json.dumps({"type": "meta", "mode": "chat", "think_enabled": thinking_enabled}) + "\n\n"

    payload = {
        "model": active_model["name"],
        "messages": messages,
        "temperature": temperature,
        "max_tokens": 4096,
        "stream": True
    }
    if context_length:
        payload["context_length"] = context_length

    try:
        resp = requests.post(LM_STUDIO_URL, json=payload, stream=True, timeout=300)
        resp.raise_for_status()

        content_buf = ""
        reasoning_buf = ""

        for line in resp.iter_lines():
            if not line:
                continue
            line = line.decode("utf-8")
            if line.startswith("data: "):
                line = line[6:]
            if line == "[DONE]":
                yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                break
            try:
                chunk = json.loads(line)
                delta = chunk["choices"][0].get("delta", {})

                if delta.get("reasoning_content"):
                    reasoning_buf += delta["reasoning_content"]
                    yield "data: " + json.dumps({"type": "reasoning", "text": delta["reasoning_content"]}) + "\n\n"

                if delta.get("content"):
                    content_buf += delta["content"]
                    yield "data: " + json.dumps({"type": "content", "text": delta["content"]}) + "\n\n"

            except (json.JSONDecodeError, KeyError, IndexError):
                continue

    except requests.exceptions.ConnectionError:
        yield "data: " + json.dumps({"type": "error", "text": "LM Studio nicht erreichbar!"}) + "\n\n"
    except Exception as e:
        yield "data: " + json.dumps({"type": "error", "text": str(e)}) + "\n\n"


@app.route("/api/models", methods=["GET"])
def list_models():
    """Nutzt den neuen /api/v1/models Endpoint der alle downloaded models zurückgibt."""
    try:
        resp = requests.get(f"{LM_API}/api/v1/models", timeout=5)
        if not resp.ok:
            return jsonify({"models": [], "active": active_model["name"], "loaded_ids": []})

        data = resp.json()
        models_raw = data.get("models", [])
    except Exception as e:
        print(f"[Models] Fehler: {e}")
        return jsonify({"models": [], "active": active_model["name"], "loaded_ids": []})

    available = []
    loaded_ids = []

    for m in models_raw:
        key = m.get("key", "")
        display_name = m.get("display_name", key)
        size_gb = round(m.get("size_bytes", 0) / 1e9, 1)
        loaded_instances = m.get("loaded_instances", [])
        is_loaded = len(loaded_instances) > 0
        is_active = active_model["name"] == key

        instance_id = loaded_instances[0]["id"] if loaded_instances else None

        if is_loaded:
            loaded_ids.append(key)

        available.append({
            "id": key,
            "name": display_name,
            "folder": m.get("publisher", ""),
            "load_id": key,
            "instance_id": instance_id,   # für unload
            "size_gb": size_gb,
            "loaded": is_loaded,
            "active": is_active,
            "type": m.get("type", "llm"),
        })

    available.sort(key=lambda x: (not x["loaded"], x["name"]))
    return jsonify({"models": available, "active": active_model["name"], "loaded_ids": loaded_ids})


@app.route("/api/models/load", methods=["POST"])
def load_model():
    data = request.json
    load_id = data.get("load_id", "")
    display_name = data.get("name", load_id)
    gpu_offload = data.get("gpu_offload", 1)  # Default: volle GPU

    if not load_id:
        return jsonify({"ok": False, "error": "Kein load_id angegeben"}), 400

    try:
        resp = requests.post(f"{LM_API}/api/v1/models/load", json={
            "model": load_id,
            "flash_attention": True,
            "offload_kv_cache_to_gpu": True
        }, timeout=60)
        if resp.ok:
            instance_id = resp.json().get("instance_id", load_id)
            active_model["name"] = instance_id
            print(f"[Model] Geladen: {display_name} (GPU: {gpu_offload*100:.0f}%)")
            return jsonify({"ok": True, "active": instance_id})
        else:
            return jsonify({"ok": False, "error": resp.text}), 400
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/models/unload", methods=["POST"])
def unload_model():
    data = request.json
    instance_id = data.get("instance_id", "")
    try:
        resp = requests.post(f"{LM_API}/api/v1/models/unload", json={"instance_id": instance_id}, timeout=30)
        if resp.ok:
            # Falls das aktive Modell entladen wird, zurücksetzen
            if active_model["name"] == instance_id:
                active_model["name"] = MODEL
            return jsonify({"ok": True})
        else:
            return jsonify({"ok": False, "error": resp.text}), 400
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/thinking/toggle", methods=["POST"])
def toggle_thinking():
    # Jetzt: UI-Toggle für Reasoning-Anzeige (kein serverseitiger Effekt mehr)
    global thinking_enabled
    thinking_enabled = not thinking_enabled
    return jsonify({"enabled": thinking_enabled})


@app.route("/api/thinking/status", methods=["GET"])
def thinking_status():
    return jsonify({"enabled": thinking_enabled})


@app.route("/api/research/toggle", methods=["POST"])
def toggle_research():
    global research_enabled
    research_enabled = not research_enabled
    print(f"[Research] Modus: {'AN' if research_enabled else 'AUS'}")
    return jsonify({"enabled": research_enabled})


@app.route("/api/research/status", methods=["GET"])
def research_status():
    return jsonify({"enabled": research_enabled})


@app.route("/api/image-generation/toggle", methods=["POST"])
def toggle_image_generation():
    global image_generation_enabled
    image_generation_enabled = not image_generation_enabled
    print(f"[ImageGen] {'AN' if image_generation_enabled else 'AUS'}")
    return jsonify({"enabled": image_generation_enabled})


@app.route("/api/image-generation/status", methods=["GET"])
def image_generation_status():
    return jsonify({"enabled": image_generation_enabled})


@app.route("/api/knowledge", methods=["GET"])
def list_knowledge():
    files = []
    allowed = {".html", ".css", ".js", ".txt", ".md", ".json"}
    for f in sorted(KNOWLEDGE_DIR.iterdir()):
        if f.is_file() and f.suffix.lower() in allowed:
            files.append({
                "name": f.name,
                "size": f.stat().st_size,
                "modified": datetime.fromtimestamp(f.stat().st_mtime).strftime("%d.%m %H:%M")
            })
    return jsonify({"files": files, "dir": str(KNOWLEDGE_DIR)})


@app.route("/api/knowledge/<filename>", methods=["GET"])
def get_knowledge_file(filename):
    path = KNOWLEDGE_DIR / Path(filename).name
    if path.exists():
        return path.read_text(encoding="utf-8"), 200, {"Content-Type": "text/plain; charset=utf-8"}
    return "Nicht gefunden", 404


@app.route("/download/<filename>")
def download(filename):
    safe = re.sub(r'[<>:"/\\|?*]', '_', filename)
    path = OUTPUT_DIR / safe
    if path.exists():
        return send_file(path, as_attachment=True)
    return "Datei nicht gefunden", 404


@app.route("/api/models/active", methods=["GET"])
def get_active_model():
    return jsonify({"active": active_model["name"]})


@app.route("/api/models/active", methods=["POST"])
def set_active_model():
    data = request.json
    model_id = data.get("id", "")
    if not model_id:
        return jsonify({"ok": False, "error": "Kein Modell-ID"}), 400
    active_model["name"] = model_id
    print(f"[Model] Aktives Modell: {model_id}")
    return jsonify({"ok": True, "active": model_id})




# ─── CHAT HISTORY ENDPOINTS ───────────────────────────────────────────────────

@app.route("/api/chats", methods=["GET"])
def list_chats():
    with get_db() as conn:
        chats = conn.execute(
            "SELECT id, title, model, created, updated FROM chats ORDER BY updated DESC"
        ).fetchall()
    return jsonify({"chats": [dict(c) for c in chats]})


@app.route("/api/chats", methods=["POST"])
def create_chat():
    now = datetime.now().isoformat()
    with get_db() as conn:
        cur = conn.execute(
            "INSERT INTO chats (title, model, created, updated) VALUES (?, ?, ?, ?)",
            ("Neuer Chat", active_model["name"], now, now)
        )
        chat_id = cur.lastrowid
    return jsonify({"id": chat_id, "title": "Neuer Chat"})


@app.route("/api/chats/<int:chat_id>", methods=["GET"])
def get_chat(chat_id):
    with get_db() as conn:
        chat = conn.execute("SELECT * FROM chats WHERE id=?", (chat_id,)).fetchone()
        if not chat:
            return jsonify({"error": "Chat nicht gefunden"}), 404
        messages = conn.execute(
            "SELECT role, content, created FROM messages WHERE chat_id=? ORDER BY id",
            (chat_id,)
        ).fetchall()
    return jsonify({
        "chat": dict(chat),
        "messages": [dict(m) for m in messages]
    })


@app.route("/api/chats/<int:chat_id>", methods=["PATCH"])
def rename_chat(chat_id):
    title = request.json.get("title", "").strip()
    if not title:
        return jsonify({"error": "Kein Titel"}), 400
    now = datetime.now().isoformat()
    with get_db() as conn:
        conn.execute("UPDATE chats SET title=?, updated=? WHERE id=?", (title, now, chat_id))
    return jsonify({"ok": True})


@app.route("/api/chats/<int:chat_id>", methods=["DELETE"])
def delete_chat(chat_id):
    with get_db() as conn:
        conn.execute("DELETE FROM chats WHERE id=?", (chat_id,))
    return jsonify({"ok": True})


@app.route("/api/chats/<int:chat_id>/messages", methods=["POST"])
def add_messages(chat_id):
    """Fügt neue Nachrichten zu einem Chat hinzu und aktualisiert Titel falls nötig."""
    data = request.json
    msgs = data.get("messages", [])  # [{role, content}]
    now = datetime.now().isoformat()

    with get_db() as conn:
        chat = conn.execute("SELECT * FROM chats WHERE id=?", (chat_id,)).fetchone()
        if not chat:
            return jsonify({"error": "Chat nicht gefunden"}), 404

        for m in msgs:
            conn.execute(
                "INSERT INTO messages (chat_id, role, content, created) VALUES (?, ?, ?, ?)",
                (chat_id, m["role"], m["content"], now)
            )

        # Auto-Titel: erste User-Nachricht (max 50 Zeichen)
        if chat["title"] == "Neuer Chat":
            first_user = next((m["content"] for m in msgs if m["role"] == "user"), None)
            if first_user:
                title = first_user[:50] + ("..." if len(first_user) > 50 else "")
                conn.execute("UPDATE chats SET title=?, updated=? WHERE id=?", (title, now, chat_id))
            else:
                conn.execute("UPDATE chats SET updated=? WHERE id=?", (now, chat_id))
        else:
            conn.execute("UPDATE chats SET updated=? WHERE id=?", (now, chat_id))

    return jsonify({"ok": True})


@app.route("/api/system-prompt", methods=["GET"])
def get_system_prompt():
    return jsonify({"prompt": custom_system_prompt if custom_system_prompt is not None else SYSTEM_PROMPT})

@app.route("/api/system-prompt", methods=["POST"])
def set_system_prompt():
    global custom_system_prompt
    data = request.json
    custom_system_prompt = data.get("prompt", "").strip() or None
    print(f"[SystemPrompt] Geändert ({len(custom_system_prompt or '')} Zeichen)")
    return jsonify({"ok": True})

@app.route("/api/system-prompt/reset", methods=["POST"])
def reset_system_prompt():
    global custom_system_prompt
    custom_system_prompt = None
    print("[SystemPrompt] Auf Default zurückgesetzt")
    return jsonify({"ok": True, "prompt": SYSTEM_PROMPT})

# ─── COMFYUI IMAGE GENERATION ─────────────────────────────────────────────────
# Workflows sind ausgelagert in workflows.py


def comfy_get_models_by_type(model_type: str) -> list:
    """Holt verfügbare Modelle von ComfyUI gefiltert nach Typ.
    Gibt immer die vollen Pfade zurück genau so wie ComfyUI sie listet.
    """
    try:
        t = model_type.lower()

        if t == "anima":
            resp = requests.get(f"{COMFY_URL}/object_info/UNETLoader", timeout=5)
            resp.raise_for_status()
            all_models = resp.json().get("UNETLoader", {}).get("input", {}).get("required", {}).get("unet_name", [None])[0] or []
            return [m for m in all_models if "anima" in m.lower()]

        elif t in ("illustrious", "zimage"):
            # Alle Checkpoints 1:1 so zurückgeben wie ComfyUI sie listet (voller Pfad)
            resp = requests.get(f"{COMFY_URL}/object_info/CheckpointLoaderSimple", timeout=5)
            resp.raise_for_status()
            all_models = resp.json().get("CheckpointLoaderSimple", {}).get("input", {}).get("required", {}).get("ckpt_name", [None])[0] or []
            print(f"[ComfyUI] Alle Checkpoints ({len(all_models)}): {all_models[:5]}")
            return all_models  # Kein Filter — User wählt selbst

    except Exception as e:
        print(f"[ComfyUI] Modelle abrufen fehlgeschlagen ({model_type}): {e}")
    return []

# Legacy-Alias für alten Code
def comfy_get_models() -> list:
    return comfy_get_models_by_type("anima")


def generate_danbooru_prompt(user_message: str) -> dict:
    """Lässt das LLM einen Danbooru-Tag-Prompt aus der User-Beschreibung generieren."""
    try:
        resp = requests.post(LM_STUDIO_URL, json={
            "model": active_model["name"],
            "messages": [
                {"role": "system", "content": DANBOORU_PROMPT},
                {"role": "user", "content": user_message}
            ],
            "temperature": 0.4,
            "max_tokens": 300,
            "stream": False
        }, timeout=30)
        resp.raise_for_status()
        text = resp.json()["choices"][0]["message"].get("content", "")
        text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL).strip()
        print(f"[Danbooru] Raw: {text[:200]}")
        parsed = parse_json_response(text)
        if parsed.get("prompt"):
            print(f"[Danbooru] ✅ {len(parsed['prompt'].split(','))} Tags generiert")
            return parsed
    except Exception as e:
        print(f"[Danbooru] Fehler: {e}")
    # Fallback
    return {
        "prompt": f"masterpiece, best quality, score_7, {user_message}",
        "negative_prompt": "worst quality, low quality, score_1, score_2, score_3, blurry",
        "aspect_ratio": "3:4 (Golden Ratio)"
    }


def comfy_generate_stream(workflow: dict, model_type: str = "anima"):
    """Generator: Sendet Workflow an ComfyUI, streamt Previews + Progress via WebSocket.
    Yieldet dicts: image_preview | image_progress | image_final | error
    """
    import struct
    try:
        import websocket
    except ImportError:
        yield {"type": "error", "text": "websocket-client nicht installiert: pip install websocket-client"}
        return

    import random
    client_id = f"office-{random.randint(10000, 99999)}"

    try:
        resp = requests.post(f"{COMFY_URL}/prompt", json={"prompt": workflow, "client_id": client_id}, timeout=10)
        resp.raise_for_status()
        prompt_id = resp.json()["prompt_id"]
        print(f"[ComfyUI] Job gestartet: {prompt_id} (type={model_type}, client={client_id})")
    except Exception as e:
        yield {"type": "error", "text": f"ComfyUI nicht erreichbar: {e}"}
        return

    ws_url = f"ws://127.0.0.1:8188/ws?clientId={client_id}"
    try:
        ws = websocket.create_connection(ws_url, timeout=300)
        print(f"[ComfyUI] WebSocket verbunden")
    except Exception as e:
        yield {"type": "error", "text": f"WebSocket fehlgeschlagen: {e}"}
        return

    try:
        while True:
            msg = ws.recv()

            # Binär = Preview-JPEG
            if isinstance(msg, bytes) and len(msg) > 8:
                event_type = struct.unpack_from(">I", msg, 0)[0]
                if event_type == 1:  # PREVIEW_IMAGE
                    b64 = base64.b64encode(msg[8:]).decode("utf-8")
                    yield {"type": "image_preview", "b64": b64}
                continue

            try:
                data = json.loads(msg)
            except Exception:
                continue

            msg_type = data.get("type", "")

            if msg_type == "progress":
                val = data["data"].get("value", 0)
                max_val = data["data"].get("max", 1)
                pct = int(val / max_val * 100) if max_val else 0
                yield {"type": "image_progress", "value": val, "max": max_val, "pct": pct}

            elif msg_type == "execution_success":
                if data["data"].get("prompt_id") == prompt_id:
                    print(f"[ComfyUI] Fertig: {prompt_id}")
                    try:
                        hist = requests.get(f"{COMFY_URL}/history/{prompt_id}", timeout=10).json()
                        outputs = hist.get(prompt_id, {}).get("outputs", {})
                        for node_out in outputs.values():
                            for img_info in node_out.get("images", []):
                                img_resp = requests.get(f"{COMFY_URL}/view", params={
                                    "filename": img_info["filename"],
                                    "type":     img_info.get("type", "output"),
                                    "subfolder": img_info.get("subfolder", "")
                                }, timeout=15)
                                img_resp.raise_for_status()
                                img_bytes = img_resp.content
                                yield {
                                    "type":      "image_final",
                                    "b64":       base64.b64encode(img_bytes).decode("utf-8"),
                                    "filename":  img_info["filename"],
                                    "img_bytes": img_bytes
                                }
                                return
                    except Exception as e:
                        yield {"type": "error", "text": f"Finales Bild laden fehlgeschlagen: {e}"}
                        return

            elif msg_type == "execution_error":
                if data["data"].get("prompt_id") == prompt_id:
                    yield {"type": "error", "text": data["data"].get("exception_message", "ComfyUI Fehler")}
                    return

    except Exception as e:
        yield {"type": "error", "text": f"WebSocket Fehler: {e}"}
    finally:
        try:
            ws.close()
        except Exception:
            pass


@app.route("/api/comfy/generate", methods=["POST"])
def api_comfy_generate():
    data = request.json
    prompt_text = data.get("prompt", "")
    aspect_ratio = data.get("aspect_ratio", "3:4 (Golden Ratio)")

    if not prompt_text:
        return jsonify({"ok": False, "error": "Kein Prompt angegeben"}), 400

    def generate():
        yield "data: " + json.dumps({"type": "step", "text": "🎨 Verbinde mit ComfyUI...", "status": "active"}) + "\n\n"

        models = comfy_get_models()
        if not models:
            yield "data: " + json.dumps({"type": "step", "text": "⚠️ Keine Modelle gefunden, nutze Standard", "status": "done"}) + "\n\n"
        else:
            short_name = models[0].split("\\")[-1].split("/")[-1]
            yield "data: " + json.dumps({"type": "step", "text": f"🤖 Modell: {short_name}", "status": "done"}) + "\n\n"

        yield "data: " + json.dumps({"type": "step", "text": "🖼️ Generiere Bild...", "status": "active"}) + "\n\n"

        result = comfy_generate(prompt_text, aspect_ratio)

        if not result["ok"]:
            yield "data: " + json.dumps({"type": "step", "text": f"❌ {result.get('error')}", "status": "error"}) + "\n\n"
            yield "data: " + json.dumps({"type": "done"}) + "\n\n"
            return

        # Bild als Base64 laden
        try:
            params = {"filename": result["filename"], "type": result["type"]}
            if result["subfolder"]:
                params["subfolder"] = result["subfolder"]
            img_resp = requests.get(f"{COMFY_URL}/view", params=params, timeout=15)
            img_resp.raise_for_status()
            b64 = base64.b64encode(img_resp.content).decode("utf-8")
            short_model = result["model"].split("\\")[-1].split("/")[-1]
            yield "data: " + json.dumps({
                "type": "image_done",
                "image_b64": b64,
                "filename": result["filename"],
                "model": short_model,
                "prompt": prompt_text
            }) + "\n\n"
        except Exception as e:
            yield "data: " + json.dumps({"type": "step", "text": f"❌ Bild laden fehlgeschlagen: {e}", "status": "error"}) + "\n\n"
            yield "data: " + json.dumps({"type": "done"}) + "\n\n"

    return Response(stream_with_context(generate()), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/api/comfy/models", methods=["GET"])
def api_comfy_models():
    return jsonify({"models": comfy_get_models()})


@app.route("/api/comfy/image-models", methods=["GET"])
def api_comfy_image_models():
    """Gibt verfügbare Modelle gefiltert nach Typ zurück."""
    model_type = request.args.get("type", "anima")
    models = comfy_get_models_by_type(model_type)
    return jsonify({"models": models, "type": model_type})




# ─── WORKFLOW EDITOR ROUTES ───────────────────────────────────────────────────

@app.route("/workflows")
def workflows_page():
    return send_from_directory("frontend", "workflows.html")


@app.route("/api/workflows", methods=["GET"])
def list_workflows():
    files = sorted(WORKFLOWS_DIR.glob("*.json"))
    result = []
    for f in files:
        try:
            wf = json.loads(f.read_text(encoding="utf-8"))
            result.append({
                "name":         f.name,
                "display_name": f.stem.replace("_", " "),
                "node_count":   len(wf),
            })
        except Exception:
            pass
    return jsonify({"workflows": result})


@app.route("/api/workflows/<path:name>", methods=["GET"])
def get_workflow(name):
    if not name.endswith(".json") or "/" in name or "\\" in name:
        return jsonify({"error": "Ungültiger Name"}), 400
    path = WORKFLOWS_DIR / name
    if not path.exists():
        return jsonify({"error": "Nicht gefunden"}), 404
    return jsonify({"workflow": json.loads(path.read_text(encoding="utf-8"))})


@app.route("/api/workflows/prompt-suggest", methods=["POST"])
def workflow_prompt_suggest():
    data = request.get_json()
    user_msg = data.get("message", "")
    if not user_msg:
        return jsonify({"error": "Keine Eingabe"}), 400
    result = generate_danbooru_prompt(user_msg)
    return jsonify({"prompt": result.get("prompt", ""), "negative_prompt": result.get("negative_prompt", "")})


@app.route("/api/workflows/run", methods=["POST"])
def run_custom_workflow():
    data = request.get_json()
    if not data or "workflow" not in data:
        return jsonify({"error": "Kein Workflow übergeben"}), 400

    def run_generator():
        yield "data: " + json.dumps({"type": "step", "text": "⚙️ Workflow wird gestartet...", "status": "active"}) + "\n\n"
        for event in comfy_generate_stream(data["workflow"]):
            etype = event.get("type")
            if etype == "image_preview":
                yield "data: " + json.dumps({"type": "image_preview", "b64": event["b64"]}) + "\n\n"
            elif etype == "image_progress":
                yield "data: " + json.dumps({"type": "image_progress", "value": event["value"], "max": event["max"], "pct": event["pct"]}) + "\n\n"
            elif etype == "image_final":
                try:
                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    save_path = EXPORT_IMG_DIR / f"{ts}_{event['filename']}"
                    save_path.write_bytes(event["img_bytes"])
                except Exception as se:
                    print(f"[WorkflowRun] Speichern fehlgeschlagen: {se}")
                yield "data: " + json.dumps({
                    "type":      "image_done",
                    "image_b64": event["b64"],
                    "filename":  event["filename"],
                }) + "\n\n"
                return
            elif etype == "error":
                yield "data: " + json.dumps({"type": "step", "text": f"❌ {event.get('text')}", "status": "error"}) + "\n\n"
                yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                return

    return Response(stream_with_context(run_generator()), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


if __name__ == "__main__":
    init_db()
    _init_active_model()
    print("\n🚀 LM Studio Office Assistant")
    print(f"   Aktives Modell: {active_model['name']}")
    print(f"   Browser öffnen: http://localhost:5000")
    print(f"   Dateien gespeichert in: {OUTPUT_DIR}\n")
    app.run(debug=False, host="0.0.0.0", port=5000)